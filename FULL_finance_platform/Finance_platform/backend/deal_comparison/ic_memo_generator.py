"""
Investment Committee Memo Generator
====================================
Purpose: Generate professional IC memos from property comparisons
Output: Word documents with executive summary, deal comparison, and recommendations
Created: November 4, 2025
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import psycopg2
from psycopg2.extras import RealDictCursor
from datetime import datetime
from typing import List, Dict, Any
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from io import BytesIO

class ICMemoGenerator:
    """Generate Investment Committee presentation memos"""
    
    def __init__(self, db_conn):
        self.conn = db_conn
    
    def generate_memo(self, comparison_id: str, output_path: str) -> str:
        """
        Generate complete IC memo for a property comparison
        
        Args:
            comparison_id: UUID of comparison set
            output_path: Path to save the Word document
            
        Returns:
            Path to generated document
        """
        # Get comparison details and deals
        comparison = self._get_comparison_details(comparison_id)
        deals = self._get_deals_with_metrics(comparison_id)
        
        # Create Word document
        doc = Document()
        
        # Set up document properties
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add content sections
        self._add_title_page(doc, comparison)
        self._add_executive_summary(doc, comparison, deals)
        self._add_deal_overview(doc, deals)
        self._add_detailed_comparison(doc, deals)
        self._add_ranking_analysis(doc, deals)
        self._add_risk_assessment(doc, deals)
        self._add_recommendation(doc, deals)
        self._add_appendix(doc, deals)
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    def _get_comparison_details(self, comparison_id: str) -> Dict:
        """Get comparison set details"""
        cur = self.conn.cursor()
        cur.execute("""
            SELECT * FROM comparison_sets 
            WHERE comparison_id = %s
        """, (comparison_id,))
        result = cur.fetchone()
        cur.close()
        return dict(result)
    
    def _get_deals_with_metrics(self, comparison_id: str) -> List[Dict]:
        """Get all deals with metrics, sorted by rank"""
        cur = self.conn.cursor()
        cur.execute("""
            SELECT 
                pd.*,
                cm.*
            FROM property_deals pd
            JOIN comparison_metrics cm ON pd.deal_id = cm.deal_id
            WHERE pd.comparison_id = %s
            ORDER BY pd.overall_rank ASC
        """, (comparison_id,))
        deals = cur.fetchall()
        cur.close()
        return [dict(deal) for deal in deals]
    
    def _add_title_page(self, doc: Document, comparison: Dict):
        """Add title page"""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run('INVESTMENT COMMITTEE MEMORANDUM')
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Comparison name
        comp_name = doc.add_paragraph()
        comp_name_run = comp_name.add_run(comparison['comparison_name'])
        comp_name_run.font.size = Pt(18)
        comp_name_run.font.bold = True
        comp_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Date and metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}\n")
        meta.add_run(f"Status: {comparison['status']}\n")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, comparison: Dict, deals: List[Dict]):
        """Add executive summary section"""
        # Section heading
        heading = doc.add_heading('Executive Summary', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Overview paragraph
        overview = doc.add_paragraph()
        overview.add_run('This memorandum presents a comprehensive analysis of ')
        overview.add_run(f"{len(deals)} real estate investment opportunities ")
        overview.add_run('under consideration. Each deal has been evaluated using standardized metrics ')
        overview.add_run('including IRR, equity multiple, cash-on-cash returns, debt coverage ratios, ')
        overview.add_run('and operational performance indicators.')
        
        doc.add_paragraph()
        
        # Top recommendation
        if deals:
            top_deal = deals[0]
            rec_heading = doc.add_heading('Top Recommendation', level=2)
            rec_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            rec_para = doc.add_paragraph()
            rec_para.add_run(f"Property: ").bold = True
            rec_para.add_run(f"{top_deal['property_name']}\n")
            
            rec_para.add_run(f"Location: ").bold = True
            rec_para.add_run(f"{top_deal['city']}, {top_deal['state']}\n")
            
            rec_para.add_run(f"Property Type: ").bold = True
            rec_para.add_run(f"{top_deal['property_type']}\n")
            
            rec_para.add_run(f"Overall Score: ").bold = True
            rec_para.add_run(f"{float(top_deal['overall_score']):.1f}/100\n")
            
            rec_para.add_run(f"Rank: ").bold = True
            rec_para.add_run(f"#{top_deal['overall_rank']} of {len(deals)}\n")
        
        doc.add_paragraph()
        
        # Key metrics summary table
        self._add_summary_metrics_table(doc, deals[:5])  # Top 5 deals
        
        doc.add_page_break()
    
    def _add_summary_metrics_table(self, doc: Document, deals: List[Dict]):
        """Add summary metrics table for top deals"""
        table = doc.add_table(rows=len(deals) + 1, cols=7)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        headers = ['Rank', 'Property', 'Type', 'Levered IRR', 'MOIC', 'CoC Y1', 'DSCR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Data rows
        for idx, deal in enumerate(deals, 1):
            row = table.rows[idx]
            
            row.cells[0].text = f"#{deal['overall_rank']}"
            row.cells[1].text = deal['property_name'][:25]  # Truncate if too long
            row.cells[2].text = deal['property_type']
            row.cells[3].text = f"{float(deal['levered_irr']) * 100:.1f}%" if deal['levered_irr'] else "N/A"
            row.cells[4].text = f"{float(deal['equity_multiple']):.2f}x" if deal['equity_multiple'] else "N/A"
            row.cells[5].text = f"{float(deal['cash_on_cash_y1']) * 100:.1f}%" if deal['cash_on_cash_y1'] else "N/A"
            row.cells[6].text = f"{float(deal['dscr_year1']):.2f}x" if deal['dscr_year1'] else "N/A"
            
            # Center align data
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_deal_overview(self, doc: Document, deals: List[Dict]):
        """Add individual deal overviews"""
        heading = doc.add_heading('Deal Summaries', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        for deal in deals:
            self._add_individual_deal_summary(doc, deal)
            doc.add_paragraph()
    
    def _add_individual_deal_summary(self, doc: Document, deal: Dict):
        """Add summary for one deal"""
        # Deal name as heading
        deal_heading = doc.add_heading(f"{deal['property_name']} (Rank #{deal['overall_rank']})", level=2)
        deal_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Property details table
        table = doc.add_table(rows=6, cols=4)
        table.style = 'Light List Accent 1'
        
        # Row 1
        table.rows[0].cells[0].text = 'Property Type'
        table.rows[0].cells[1].text = deal['property_type']
        table.rows[0].cells[2].text = 'Location'
        table.rows[0].cells[3].text = f"{deal['city']}, {deal['state']}" if deal['city'] else "N/A"
        
        # Row 2
        table.rows[1].cells[0].text = 'Total SF'
        table.rows[1].cells[1].text = f"{int(deal['total_sf']):,}" if deal['total_sf'] else "N/A"
        table.rows[1].cells[2].text = 'Units/Rooms'
        table.rows[1].cells[3].text = f"{int(deal['units']):,}" if deal['units'] else "N/A"
        
        # Row 3
        table.rows[2].cells[0].text = 'Levered IRR'
        table.rows[2].cells[1].text = f"{float(deal['levered_irr']) * 100:.1f}%" if deal['levered_irr'] else "N/A"
        table.rows[2].cells[2].text = 'Equity Multiple'
        table.rows[2].cells[3].text = f"{float(deal['equity_multiple']):.2f}x" if deal['equity_multiple'] else "N/A"
        
        # Row 4
        table.rows[3].cells[0].text = 'CoC Year 1'
        table.rows[3].cells[1].text = f"{float(deal['cash_on_cash_y1']) * 100:.1f}%" if deal['cash_on_cash_y1'] else "N/A"
        table.rows[3].cells[2].text = 'DSCR Year 1'
        table.rows[3].cells[3].text = f"{float(deal['dscr_year1']):.2f}x" if deal['dscr_year1'] else "N/A"
        
        # Row 5
        table.rows[4].cells[0].text = 'Equity Required'
        table.rows[4].cells[1].text = f"${float(deal['equity_required']) / 1e6:.1f}M" if deal['equity_required'] else "N/A"
        table.rows[4].cells[2].text = 'Purchase Price'
        table.rows[4].cells[3].text = f"${float(deal['purchase_price']) / 1e6:.1f}M" if deal['purchase_price'] else "N/A"
        
        # Row 6
        table.rows[5].cells[0].text = 'Overall Score'
        table.rows[5].cells[1].text = f"{float(deal['overall_score']):.1f}/100" if deal['overall_score'] else "N/A"
        table.rows[5].cells[2].text = 'Risk Score'
        table.rows[5].cells[3].text = f"{float(deal['risk_score']):.1f}/100" if deal['risk_score'] else "N/A"
        
        # Bold the labels
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[2].paragraphs[0].runs[0].font.bold = True
    
    def _add_detailed_comparison(self, doc: Document, deals: List[Dict]):
        """Add detailed side-by-side comparison table"""
        doc.add_page_break()
        
        heading = doc.add_heading('Detailed Comparison Matrix', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Create comparison table
        table = doc.add_table(rows=len(deals) + 1, cols=10)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = ['Rank', 'Property', 'Type', 'IRR', 'MOIC', 'CoC Y1', 'DSCR', 
                   'Entry Cap', 'NOI Margin', 'Score']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows with conditional formatting
        for idx, deal in enumerate(deals, 1):
            row = table.rows[idx]
            
            row.cells[0].text = f"#{deal['overall_rank']}"
            row.cells[1].text = deal['property_name'][:20]
            row.cells[2].text = deal['property_type'][:15]
            row.cells[3].text = f"{float(deal['levered_irr']) * 100:.1f}%" if deal['levered_irr'] else "N/A"
            row.cells[4].text = f"{float(deal['equity_multiple']):.2f}x" if deal['equity_multiple'] else "N/A"
            row.cells[5].text = f"{float(deal['cash_on_cash_y1']) * 100:.1f}%" if deal['cash_on_cash_y1'] else "N/A"
            row.cells[6].text = f"{float(deal['dscr_year1']):.2f}x" if deal['dscr_year1'] else "N/A"
            row.cells[7].text = f"{float(deal['entry_cap_rate']) * 100:.1f}%" if deal['entry_cap_rate'] else "N/A"
            row.cells[8].text = f"{float(deal['noi_margin']) * 100:.0f}%" if deal['noi_margin'] else "N/A"
            row.cells[9].text = f"{float(deal['overall_score']):.0f}" if deal['overall_score'] else "N/A"
            
            # Color coding based on score
            score = float(deal['overall_score']) if deal['overall_score'] else 0
            if score >= 80:
                color = RGBColor(0, 176, 80)  # Green
            elif score >= 60:
                color = RGBColor(255, 192, 0)  # Yellow
            else:
                color = RGBColor(255, 0, 0)  # Red
            
            row.cells[9].paragraphs[0].runs[0].font.color.rgb = color
            row.cells[9].paragraphs[0].runs[0].font.bold = True
    
    def _add_ranking_analysis(self, doc: Document, deals: List[Dict]):
        """Add ranking analysis with charts"""
        doc.add_page_break()
        
        heading = doc.add_heading('Ranking Analysis', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Top performers by category
        doc.add_heading('Top Performers by Category', level=2)
        
        # Returns
        sorted_by_irr = sorted(deals, key=lambda x: float(x['levered_irr'] or 0), reverse=True)[:3]
        para = doc.add_paragraph()
        para.add_run('Highest IRR: ').bold = True
        para.add_run(', '.join([f"{d['property_name']} ({float(d['levered_irr']) * 100:.1f}%)" 
                                for d in sorted_by_irr if d['levered_irr']]))
        
        # Risk
        sorted_by_dscr = sorted(deals, key=lambda x: float(x['dscr_year1'] or 0), reverse=True)[:3]
        para = doc.add_paragraph()
        para.add_run('Best Debt Coverage: ').bold = True
        para.add_run(', '.join([f"{d['property_name']} ({float(d['dscr_year1']):.2f}x)" 
                                for d in sorted_by_dscr if d['dscr_year1']]))
        
        # Operations
        sorted_by_margin = sorted(deals, key=lambda x: float(x['noi_margin'] or 0), reverse=True)[:3]
        para = doc.add_paragraph()
        para.add_run('Highest NOI Margin: ').bold = True
        para.add_run(', '.join([f"{d['property_name']} ({float(d['noi_margin']) * 100:.0f}%)" 
                                for d in sorted_by_margin if d['noi_margin']]))
    
    def _add_risk_assessment(self, doc: Document, deals: List[Dict]):
        """Add risk assessment section"""
        doc.add_page_break()
        
        heading = doc.add_heading('Risk Assessment', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Risk tiers
        doc.add_heading('Risk-Adjusted Rankings', level=2)
        
        # Sort by risk-adjusted rank
        for deal in sorted(deals, key=lambda x: x['risk_adjusted_rank'] or 999)[:5]:
            para = doc.add_paragraph()
            para.add_run(f"#{deal['risk_adjusted_rank']} {deal['property_name']}: ").bold = True
            
            risk_score = float(deal['risk_score']) if deal['risk_score'] else 0
            if risk_score >= 75:
                risk_level = "LOW RISK"
                color = RGBColor(0, 176, 80)
            elif risk_score >= 50:
                risk_level = "MODERATE RISK"
                color = RGBColor(255, 192, 0)
            else:
                risk_level = "HIGH RISK"
                color = RGBColor(255, 0, 0)
            
            run = para.add_run(f" {risk_level}")
            run.font.color.rgb = color
            run.bold = True
            
            para.add_run(f" (Risk Score: {risk_score:.0f}/100)")
    
    def _add_recommendation(self, doc: Document, deals: List[Dict]):
        """Add final recommendation section"""
        doc.add_page_break()
        
        heading = doc.add_heading('Investment Committee Recommendation', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        if not deals:
            doc.add_paragraph('No deals to recommend.')
            return
        
        top_deal = deals[0]
        
        # Recommendation
        rec_heading = doc.add_heading('RECOMMEND APPROVAL', level=2)
        rec_heading.runs[0].font.color.rgb = RGBColor(0, 176, 80)
        
        para = doc.add_paragraph()
        para.add_run('Property: ').bold = True
        para.add_run(f"{top_deal['property_name']}\n")
        
        para.add_run('Rationale: ').bold = True
        para.add_run(f"This property ranks #{top_deal['overall_rank']} out of {len(deals)} opportunities ")
        para.add_run(f"analyzed, with an overall score of {float(top_deal['overall_score']):.1f}/100. ")
        para.add_run(f"The deal offers a levered IRR of {float(top_deal['levered_irr']) * 100:.1f}%, ")
        para.add_run(f"equity multiple of {float(top_deal['equity_multiple']):.2f}x, ")
        para.add_run(f"and strong debt coverage with a DSCR of {float(top_deal['dscr_year1']):.2f}x.")
        
        doc.add_paragraph()
        
        # Investment highlights
        doc.add_heading('Key Investment Highlights', level=3)
        highlights = doc.add_paragraph(style='List Bullet')
        highlights.add_run(f"Superior returns with {float(top_deal['levered_irr']) * 100:.1f}% IRR")
        
        highlights = doc.add_paragraph(style='List Bullet')
        highlights.add_run(f"Strong risk profile (Risk Score: {float(top_deal['risk_score']):.0f}/100)")
        
        highlights = doc.add_paragraph(style='List Bullet')
        highlights.add_run(f"Excellent market location in {top_deal['city']}, {top_deal['state']}")
        
        # Next steps
        doc.add_heading('Recommended Next Steps', level=3)
        steps = doc.add_paragraph(style='List Number')
        steps.add_run('Approve preliminary investment thesis')
        
        steps = doc.add_paragraph(style='List Number')
        steps.add_run('Authorize full due diligence process')
        
        steps = doc.add_paragraph(style='List Number')
        steps.add_run('Engage legal counsel for transaction documents')
        
        steps = doc.add_paragraph(style='List Number')
        steps.add_run('Schedule site visit and management meetings')
    
    def _add_appendix(self, doc: Document, deals: List[Dict]):
        """Add appendix with additional details"""
        doc.add_page_break()
        
        heading = doc.add_heading('Appendix: Detailed Metrics', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Full metrics table for all deals
        table = doc.add_table(rows=len(deals) + 1, cols=8)
        table.style = 'Light Grid'
        
        headers = ['Property', 'Equity Req', 'NOI Y1', 'Exit Value', 
                   'Unlevered IRR', 'LTV', 'Debt Yield', 'Cap Rate']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        for idx, deal in enumerate(deals, 1):
            row = table.rows[idx]
            row.cells[0].text = deal['property_name'][:25]
            row.cells[1].text = f"${float(deal['equity_required']) / 1e6:.1f}M" if deal['equity_required'] else "N/A"
            row.cells[2].text = f"${float(deal['noi_year1']) / 1e6:.2f}M" if deal['noi_year1'] else "N/A"
            row.cells[3].text = f"${float(deal['gross_exit_value']) / 1e6:.1f}M" if deal.get('gross_exit_value') else "N/A"
            row.cells[4].text = f"{float(deal['unlevered_irr']) * 100:.1f}%" if deal['unlevered_irr'] else "N/A"
            row.cells[5].text = f"{float(deal['ltv']) * 100:.0f}%" if deal['ltv'] else "N/A"
            row.cells[6].text = f"{float(deal['debt_yield']) * 100:.1f}%" if deal.get('debt_yield') else "N/A"
            row.cells[7].text = f"{float(deal['entry_cap_rate']) * 100:.1f}%" if deal['entry_cap_rate'] else "N/A"

# =====================================================
# API ENDPOINT FOR IC MEMO GENERATION
# =====================================================

def generate_ic_memo_api(comparison_id: str, db_conn) -> str:
    """
    API endpoint to generate IC memo
    
    Args:
        comparison_id: UUID of comparison set
        db_conn: Database connection
        
    Returns:
        Path to generated document
    """
    output_path = f"/tmp/IC_Memo_{comparison_id}.docx"
    
    generator = ICMemoGenerator(db_conn)
    generator.generate_memo(comparison_id, output_path)
    
    return output_path

if __name__ == "__main__":
    # Example usage
    import psycopg2
    
    conn = psycopg2.connect(
        host="localhost",
        database="portfolio_dashboard",
        user="postgres",
        password="your_password",
        cursor_factory=psycopg2.extras.RealDictCursor
    )
    
    # Generate memo for a comparison
    comparison_id = "your-comparison-id-here"
    output_path = generate_ic_memo_api(comparison_id, conn)
    print(f"IC Memo generated: {output_path}")
    
    conn.close()
